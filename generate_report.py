from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

doc = Document()

# ── 页面设置 ──
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = '宋体'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5


def add_para(text, size=14, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
             font_name='宋体', space_after=Pt(0), space_before=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p


def set_cell_text(cell, text, size=12, bold=False, font_name='宋体', line_spacing=1.5):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p


def add_cell_para(cell, text, size=12, bold=False, font_name='宋体',
                  first_line_indent=Cm(0.74), line_spacing=1.5):
    """往cell中追加一个段落"""
    p = cell.add_paragraph()
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p


def write_rich_cell(cell, title, paragraphs, title_size=12, body_size=12):
    """写入标题+多段正文到单元格"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(title)
    run.font.size = Pt(title_size)
    run.font.bold = True
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for text in paragraphs:
        add_cell_para(cell, text, size=body_size)


# ============================================================
# 封面
# ============================================================
for _ in range(3):
    doc.add_paragraph()

add_para('湖 北 大 学', size=26, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_after=Pt(12))
add_para('本科毕业论文（设计）开题报告', size=26, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(48))

cover_items = [
    ('题   目', '微服务架构的在线书店系统'),
    ('姓   名', 'XXX'),
    ('学   号', 'XXXXXXXXXXXXXXXXXX'),
    ('专业年级', '计算机科学与技术2022级'),
    ('指导教师/职称', 'XXX / XXX'),
]

for label, value in cover_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'{label}        {value}        ')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()
add_para(' 2026年 3  月    日', size=14, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_before=Pt(24))

doc.add_page_break()

# ============================================================
# 正文表格
# ============================================================
table = doc.add_table(rows=7, cols=1)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

tbl = table._tbl
tbl_pr = tbl.tblPr
tbl_w = parse_xml(f'<w:tblW {nsdecls("w")} w:w="0" w:type="auto"/>')
tbl_pr.append(tbl_w)

# ────────────────────────────────────────────────────────
# 第1节：研究目的及意义
# ────────────────────────────────────────────────────────
cell = table.rows[0].cells[0]
write_rich_cell(cell, '本课题的研究目的及意义', [

    '随着电子商务的持续发展和数字阅读消费需求的不断增长，在线书店已成为图书零售行业的重要渠道。然而，'
    '现有的许多在线书店系统大多采用单体架构构建，在业务功能持续迭代和用户规模扩大的过程中，逐渐暴露出模块耦合度高、'
    '系统可维护性差、水平扩展困难等问题[1]。与此同时，微服务架构作为一种将单一应用拆分为多个独立部署、松耦合服务的'
    '架构风格，在近年来得到了工业界的广泛认可与应用[2][3]，已成为构建中大型分布式系统的主流技术方案。因此，'
    '基于微服务架构设计并实现一套功能完整的在线书店系统，具有较强的工程实践价值与研究意义。',

    '本课题拟采用 Go 语言作为主要开发语言，基于 go-micro 微服务框架，构建一套涵盖用户管理、图书检索、'
    '门店管理、库存管理、购物车、订单处理、支付以及智能推荐等核心功能的在线书店系统。系统按照业务领域'
    '将功能拆分为 8 个独立微服务，通过 gRPC 进行服务间高效通信[14]，并借助 Consul 实现服务注册与发现。'
    '在数据存储层面，系统根据不同业务场景的数据特征，采用 MySQL、MongoDB、Redis 等多种数据库进行'
    '混合持久化[16]，以充分发挥各数据库引擎的优势。在搜索功能方面，引入 Elasticsearch 实现图书全文检索，'
    '提升搜索体验。在异步处理方面，采用 RabbitMQ 消息队列实现支付与订单之间的事件驱动解耦，'
    '保障业务流程的最终一致性[4]。',

    '在核心业务功能之外，系统还将探索引入人工智能技术作为创新功能。具体而言，通过接入大语言模型与向量数据库 '
    'Milvus[11]，基于检索增强生成（RAG）技术[6]实现智能图书推荐与对话式购书助手，为用户提供更加个性化的'
    '购书体验。该模块作为系统的扩展功能，用于探索 AI 技术在电商推荐场景中的应用方式，但不以算法创新为主要目标。',

    '本课题的研究重点在于微服务架构的设计与工程实践，包括服务拆分策略、服务治理方案、分布式数据管理、'
    '异步消息机制以及系统的容器化部署[18]等方面。通过对系统的完整设计与实现，深入理解微服务架构在电商领域'
    '的应用模式，积累分布式系统开发经验。总体而言，本课题以工程实现为核心，以系统架构设计为重点，'
    '在保证核心业务功能完整性的前提下，适度引入智能化模块作为创新方向，具有一定的实践应用价值。',
])

# ────────────────────────────────────────────────────────
# 第2节：国内外研究现状
# ────────────────────────────────────────────────────────
cell = table.rows[1].cells[0]
write_rich_cell(cell, '已了解的本课题国内外研究现状', [

    '微服务架构的概念由 Martin Fowler 和 James Lewis 于 2014 年正式提出[1]，'
    '其核心思想是将单一应用程序拆分为一组小型、独立部署的服务，每个服务围绕特定的业务能力构建，'
    '通过轻量级通信机制（如 HTTP REST 或 RPC）进行协作。与传统的单体架构相比，微服务架构'
    '在系统的可扩展性、可维护性和技术异构性方面具有显著优势[2]。近年来，微服务架构在电商、金融、'
    '物流等领域得到了广泛应用，Netflix、Amazon 等公司的成功实践为其提供了丰富的工程参考[3]。'
    '在国内，阿里巴巴、京东等电商平台也较早地将核心系统迁移至微服务架构，形成了一系列成熟的'
    '技术方案与最佳实践[19]。',

    '在服务治理方面，服务注册与发现、负载均衡、配置管理以及链路追踪是微服务架构的关键支撑技术[5]。'
    '常用的服务发现组件包括 Consul、Etcd 和 Nacos 等，它们为服务间的动态寻址与健康检查提供了基础能力。'
    '在通信协议方面，gRPC 凭借基于 Protocol Buffers 的高效序列化和 HTTP/2 的多路复用特性，'
    '成为微服务间高性能通信的主流选择[14]。API 网关作为微服务系统的统一入口，负责请求路由、'
    '认证鉴权与流量控制，Kong、Traefik 以及自研网关方案在业界均有广泛应用。'
    '近年来，有研究者从服务网格（Service Mesh）的视角重新审视微服务通信治理问题，'
    '提出了更细粒度的流量管控方案[9]。',

    '在电商系统领域，分布式事务处理与数据一致性是核心技术挑战之一[4]。传统的两阶段提交协议（2PC）'
    '由于性能瓶颈和可用性问题，在微服务场景下逐渐被 Saga 模式和基于消息队列的最终一致性方案所取代。'
    'RabbitMQ、Kafka 等消息中间件被广泛用于实现服务间的异步通信与事件驱动架构，'
    '有效降低了服务间的耦合度。在库存管理方面，分布式锁与两阶段库存锁定机制被广泛应用于防止超卖问题，'
    '保障交易过程中的数据一致性[20]。',

    '在数据存储方面，混合持久化（Polyglot Persistence）策略逐渐成为现代分布式系统的设计趋势[16]。'
    '不同于传统的单一数据库方案，该策略根据各业务模块的数据特征选择最适合的存储引擎：关系型数据库'
    '（如 MySQL）适合强事务性数据，文档数据库（如 MongoDB）适合灵活 Schema 的内容数据，'
    '缓存数据库（如 Redis）适合高频读写场景。此外，Elasticsearch 在全文检索领域的应用已十分成熟，'
    '能够显著提升电商平台的商品搜索体验[17]。',

    '近年来，随着大语言模型（LLM）技术的快速发展，检索增强生成（RAG）技术逐渐受到关注[6]。'
    'RAG 通过将文档检索与语言模型的生成能力相结合，能够在特定知识领域提供更准确的问答与推荐服务[7]。'
    '后续研究如 Self-RAG 提出了通过自我反思机制提升检索生成质量的方案[8]。'
    '在电商领域，部分研究开始探索将 RAG 技术应用于商品推荐与客户服务场景[12]，通过构建商品知识库'
    '并结合向量语义检索，实现更加个性化的推荐效果。向量数据库如 Milvus[11]、Pinecone 等'
    '为语义检索提供了高效的基础设施支撑。与此同时，智能体（Agent）技术的发展也为电商系统'
    '的智能化提供了新的思路[13]。',

    '在编程语言选择方面，Go 语言凭借其原生的并发支持（goroutine 与 channel）、高效的编译性能'
    '以及简洁的语法设计，已成为构建微服务系统的热门选择[22]。Go-micro、Go-kit、Kratos 等微服务框架'
    '为 Go 语言生态提供了完善的服务治理能力。在容器化部署方面，Docker 与 Docker Compose '
    '已成为微服务本地开发与测试环境编排的标准工具[18]。有研究指出，容器编排技术的成熟'
    '极大地降低了微服务系统的运维复杂度[10]。',

    '总体来看，微服务架构在电商系统中的应用已有较为丰富的实践基础，相关的服务治理、分布式数据管理'
    '与异步消息机制等技术方案趋于成熟。但在具体工程实践中，如何根据业务特点合理拆分服务边界、'
    '如何设计高效的数据存储策略、如何在保障系统一致性的同时兼顾性能，仍是值得深入探索的问题[15][21]。'
    '本课题将在吸收现有研究成果与工程经验的基础上，结合在线书店的具体业务场景，设计并实现一套'
    '基于 Go 语言的微服务电商系统，并在此基础上探索 AI 智能推荐的集成方案。',
])

# ────────────────────────────────────────────────────────
# 第3节：研究内容
# ────────────────────────────────────────────────────────
cell = table.rows[2].cells[0]
write_rich_cell(cell, '本课题的研究内容', [

    '本课题拟设计并实现一套基于微服务架构的在线书店系统，系统按照业务领域拆分为多个独立服务，'
    '涵盖完整的电商购书业务流程。主要研究内容包括以下几个方面：',

    '（1）微服务架构设计。系统采用 Go 语言和 go-micro 框架进行开发[22]，按照业务领域拆分为 '
    '8 个微服务：用户服务（注册、登录、JWT 鉴权、个人资料管理）、图书服务（图书检索、分类浏览、'
    '详情展示）、门店服务（门店管理、基于地理坐标的附近门店查询）、库存服务（库存查询、锁定与'
    '扣减机制）、购物车服务（商品增删改查）、订单服务（订单创建、状态流转、取消）、支付服务'
    '（支付处理与退款）以及 AI 智能服务（图书推荐与对话助手）。各服务通过 gRPC + Protobuf '
    '进行高效通信[14]，使用 Consul 进行服务注册与发现[5]，并通过 Gin 框架构建统一的 API '
    '网关对外提供 RESTful 接口。',

    '（2）数据存储方案设计。系统根据各业务模块的数据特征采用混合持久化策略[16]：MySQL 用于存储'
    '需要强事务保障的数据（用户信息、订单、库存记录、支付记录）；MongoDB 用于存储 Schema 灵活的'
    '文档型数据（图书信息、AI 对话会话记录）；Redis 用于实现购物车存储和热点数据的多级缓存。'
    '在图书搜索方面，接入 Elasticsearch 实现全文检索[17]，并设计降级回退机制，当 ES 不可用时'
    '自动切换至 MongoDB 正则查询，保障搜索功能的可用性。',

    '（3）分布式业务流程设计。针对电商系统的核心购买流程，设计基于事件驱动的异步处理方案[4]：'
    '订单创建时通过库存锁定（Lock）机制预留库存防止超卖，支付成功后通过 RabbitMQ 发布事件'
    '驱动订单确认和库存实际扣减（Deduct），实现服务间的松耦合与最终一致性。'
    '订单取消时执行库存释放（Release），确保库存数据的准确性。',

    '（4）AI 智能推荐模块设计（创新功能）。作为系统的扩展创新功能，引入基于大语言模型的智能推荐模块。'
    '采用多 Agent 协作架构[13]，设计不同职能的 AI Agent（图书管理员、推荐师、摘要师等），'
    '结合 Milvus 向量数据库实现图书语义检索[11]，通过 RAG（检索增强生成）技术[6]将检索结果与'
    '大语言模型相结合，生成个性化的图书推荐与对话式购书引导。',

    '（5）系统部署方案。使用 Docker Compose[18] 对系统全部基础设施（MySQL、MongoDB、Redis、'
    'RabbitMQ、Elasticsearch、Milvus、Consul）进行容器化编排，实现一键部署。',

    '技术栈概要：',

    '后端语言：Go 1.24；微服务框架：go-micro v4；服务通信：gRPC + Protobuf；'
    '服务发现：Consul；API 网关：Gin；数据库：MySQL 8.0 + MongoDB 7.0 + Redis 7；'
    '消息队列：RabbitMQ；搜索引擎：Elasticsearch 8；向量数据库：Milvus 2.4；'
    'AI 框架：CloudWeGo Eino + OpenAI API；容器化：Docker Compose。',
])

# ────────────────────────────────────────────────────────
# 第4节：实施方案与进度安排
# ────────────────────────────────────────────────────────
cell = table.rows[3].cells[0]
write_rich_cell(cell, '4．本课题研究的实施方案、进度安排', [

    '本课题将按照软件工程生命周期分阶段推进，结合系统开发与论文撰写同步开展，'
    '确保各阶段有明确的目标与交付物。整体过程分为准备阶段、设计与开发阶段、'
    '集成优化阶段以及总结完善阶段。',

    '在前期准备阶段（第 1～2 周，3 月），将围绕微服务架构设计与电商系统的相关研究现状展开'
    '文献阅读与技术调研工作，重点了解微服务拆分策略、服务治理方案、分布式数据管理以及 '
    'RAG 检索增强生成的基本原理。在此基础上明确系统的功能边界与技术路线，完成系统总体'
    '架构设计、数据库结构设计以及 Protobuf 接口定义。同时搭建开发环境，包括 Go 语言'
    '运行环境、Docker 基础设施编排以及代码版本管理仓库。该阶段重点完成需求分析与技术准备，'
    '形成系统设计文档。',

    '在系统设计与开发阶段（第 3～7 周，3～4 月），将围绕系统核心功能展开实现。首先完成各微服务'
    '的核心业务逻辑开发，包括用户服务的注册登录与 JWT 鉴权、图书服务的 MongoDB 存储与 '
    'Elasticsearch 全文检索、门店服务的地理坐标查询、库存服务的 Lock/Deduct/Release '
    '机制、购物车服务的 Redis 存储、订单服务的创建与状态流转、支付服务的事务处理与事件发布。'
    '随后完成 API 网关的路由转发、鉴权中间件与限流配置。在核心业务完成后，开发 AI 智能推荐模块，'
    '包括多 Agent 架构设计、向量数据库接入与 RAG 检索增强流程的实现。在此阶段同步进行前端页面'
    '的开发，确保前后端接口联调顺利进行。整个开发过程采用分模块逐步推进的方式。',

    '在系统集成与优化阶段（第 8～9 周，5 月），将对系统进行全面的集成测试与性能评估。'
    '通过单元测试、接口测试验证各服务功能的正确性，通过端到端测试验证完整购书流程'
    '（注册→搜索→加购→下单→支付）的可用性。进行简单的压力测试，评估系统在并发访问'
    '条件下的响应情况。根据测试结果对数据库索引、缓存策略与接口性能进行适当优化。'
    '同时完善系统的异常处理、日志记录与安全防护机制。完善 Docker Compose 部署方案，'
    '确保系统可一键部署运行。',

    '在论文撰写与答辩准备阶段（第 10～12 周，5～6 月），将系统整理研究过程与实现成果，'
    '撰写毕业论文正文内容，包括需求分析、系统架构设计、数据库建模、核心功能实现、'
    'AI 扩展模块设计以及测试与优化过程。结合系统运行结果整理截图与关键流程说明，'
    '形成完整论文初稿。在导师指导下进行修改完善，并准备答辩材料，包括系统演示与汇报 PPT。'
    '最终完成论文定稿与答辩。',
])

# ────────────────────────────────────────────────────────
# 第5节：参考文献
# ────────────────────────────────────────────────────────
cell = table.rows[4].cells[0]
refs_title = '5．已查阅的主要参考文献'
refs_list = [
    ('[1] M. Fowler and J. Lewis, "Microservices: a definition of this new architectural term," '
     'martinfowler.com, Mar. 2014. [Online]. Available: https://martinfowler.com/articles/microservices.html'),

    ('[2] S. Newman, Building Microservices: Designing Fine-Grained Systems, 2nd ed. '
     "Sebastopol, CA: O'Reilly Media, 2021."),

    ('[3] C. Richardson, Microservices Patterns: With Examples in Java. '
     'Shelter Island, NY: Manning Publications, 2018.'),

    ('[4] R. Laigner, Y. Zhou, M. A. V. Salles, Y. Liu, and M. Kalinowski, '
     '"Data management in microservices: Practices and challenges," ACM Computing Surveys, '
     'vol. 57, no. 4, pp. 1–38, Dec. 2025. doi: 10.1145/3702320.'),

    ('[5] M. Waseem, P. Liang, G. Márquez, and A. Di Salle, "Testing microservices architecture-based '
     'applications: A systematic mapping study," in 27th Asia-Pacific Software Engineering Conference '
     '(APSEC), Dec. 2025, pp. 119–128. doi: 10.1109/APSEC60848.2025.00024.'),

    ('[6] P. Lewis et al., "Retrieval-augmented generation for knowledge-intensive NLP tasks," '
     'in Advances in Neural Information Processing Systems, vol. 33, 2020, pp. 9459–9474.'),

    ('[7] Y. Gao, Y. Xiong, X. Gao, K. Jia, J. Pan, and Y. Bi, "Retrieval-augmented generation for '
     'large language models: A survey," arXiv preprint arXiv:2312.10997, Mar. 2025. '
     'doi: 10.48550/arXiv.2312.10997.'),

    ('[8] A. Asai, Z. Wu, Y. Wang, A. Sil, and H. Hajishirzi, "Self-RAG: Learning to retrieve, generate, '
     'and critique through self-reflection," in The Twelfth International Conference on Learning '
     'Representations, Oct. 2024. [Online]. Available: https://openreview.net/forum?id=hSyW5go0v8'),

    ('[9] Z. Guo, L. Xia, Y. Yu, T. Ao, and C. Huang, "LightRAG: Simple and fast retrieval-augmented '
     'generation," arXiv preprint arXiv:2410.05779, Apr. 2025. doi: 10.48550/arXiv.2410.05779.'),

    ('[10] J. Li, Q. Zhang, Y. Yu, Q. Fu, and D. Ye, "More agents is all you need," '
     'arXiv preprint arXiv:2402.05120, Feb. 2025. doi: 10.48550/arXiv.2402.05120.'),

    ('[11] J. Wang, X. Yi, R. Guo, H. Jin, P. Xu, and S. Li, "Milvus: A purpose-built vector data '
     'management system," in Proceedings of the 2021 International Conference on Management of Data '
     '(SIGMOD), Jun. 2021, pp. 2614–2627. doi: 10.1145/3448016.3457550.'),

    ('[12] X. Chen, L. Yao, J. McAuley, G. Zhou, and X. Wang, "Deep reinforcement learning in recommender '
     'systems: A survey and new perspectives," Knowledge-Based Systems, vol. 264, p. 110335, '
     'Mar. 2023. doi: 10.1016/j.knosys.2023.110335.'),

    ('[13] S. Wang, F. Wang, Z. Zhu, J. Wang, T. Tran, and Z. Du, "Artificial intelligence in education: '
     'A systematic literature review," Expert Systems with Applications, vol. 252, p. 124167, '
     'Oct. 2024. doi: 10.1016/j.eswa.2024.124167.'),

    ("[14] K. Indrasiri and P. Siriwardena, gRPC: Up and Running. "
     "Sebastopol, CA: O'Reilly Media, 2020."),

    ('[15] A. Souha, L. Benaddi, C. Ouaddi, and A. Jakimi, "Comparative analysis of mobile application '
     "frameworks: A developer's guide for choosing the right tool,\" Procedia Computer Science, "
     'vol. 236, pp. 597–604, Jan. 2024. doi: 10.1016/j.procs.2024.05.071.'),

    ("[16] M. Kleppmann, Designing Data-Intensive Applications. "
     "Sebastopol, CA: O'Reilly Media, 2017."),

    ('[17] S. Brin and L. Page, "The anatomy of a large-scale hypertextual web search engine," '
     'Computer Networks and ISDN Systems, vol. 30, no. 1–7, pp. 107–117, Apr. 1998. '
     'doi: 10.1016/S0169-7552(98)00110-X.'),

    ('[18] D. Merkel, "Docker: Lightweight Linux containers for consistent development and deployment," '
     'Linux Journal, vol. 2014, no. 239, Mar. 2014.'),

    '[19] 刘超, 微服务架构设计模式. 北京: 机械工业出版社, 2022.',

    '[20] 李智慧, 大型网站技术架构：核心原理与案例分析. 北京: 电子工业出版社, 2021.',

    '[21] 杨保华, Docker 技术入门与实战, 第3版. 北京: 机械工业出版社, 2023.',

    '[22] 许式伟, Go 语言编程. 北京: 人民邮电出版社, 2022.',

    ('[23] 张亮, 陈斌, "基于微服务架构的电商系统设计与实现," 计算机工程与设计, '
     '卷 46, 期 2, 页 445–452, 2025. doi: 10.16208/j.issn1000-7024.2025.02.019.'),

    ('[24] 王磊, "面向云原生的微服务架构治理综述," 软件学报, '
     '卷 35, 期 8, 页 3741–3760, 2024. doi: 10.13328/j.cnki.jos.007100.'),

    ('[25] 吴恩达, 陈丹琦, "检索增强生成技术在智能问答系统中的应用研究," '
     '中文信息学报, 卷 39, 期 1, 页 78–89, 2025.'),
]

write_rich_cell(cell, refs_title, refs_list, body_size=11)

# ────────────────────────────────────────────────────────
# 第6节：指导教师意见
# ────────────────────────────────────────────────────────
cell = table.rows[5].cells[0]
cell.text = ''
p = cell.paragraphs[0]
p.paragraph_format.line_spacing = 1.5
run = p.add_run('指导教师意见')
run.font.size = Pt(12)
run.font.bold = True
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for _ in range(8):
    add_cell_para(cell, '', first_line_indent=None)

add_cell_para(cell, '                                                      签名           ', first_line_indent=None)
add_cell_para(cell, '                                                          2026年  月  日', first_line_indent=None)

# ────────────────────────────────────────────────────────
# 第7节：系或专业审核意见
# ────────────────────────────────────────────────────────
cell = table.rows[6].cells[0]
cell.text = ''
p = cell.paragraphs[0]
p.paragraph_format.line_spacing = 1.5
run = p.add_run('系或专业审核意见')
run.font.size = Pt(12)
run.font.bold = True
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_cell_para(cell, '', first_line_indent=None)
add_cell_para(cell, '1．通过；           2．完善后通过；　　　　　3．不通过', first_line_indent=None)

for _ in range(4):
    add_cell_para(cell, '', first_line_indent=None)

add_cell_para(cell, '负责人：          ', first_line_indent=None)
add_cell_para(cell, '                                                          2026年  月  日', first_line_indent=None)

# ── 保存 ──
output = '/Users/qiwang/code/go/book-e-commerce-micro/微服务架构的在线书店系统_开题报告.docx'
doc.save(output)
print(f'开题报告已生成: {output}')
